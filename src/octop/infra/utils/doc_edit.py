"""Editable-document converters (binary document <-> Markdown) for the workspace.

The workspace online editor renders an editable document as Markdown in Monaco
and saves back through ``PUT /workspace/doc``. Every supported extension
registers a :class:`DocConverter` here; adding a new editable extension is one
converter class plus a matching entry in the dashboard registry — the API
endpoints, viewer and save routing are generic and never change.

Currently supported: ``.docx``. The Markdown round-trip keeps headings,
bold/italic/inline code, bullet & numbered lists, pipe tables, blockquotes and
fenced code blocks. Complex formatting (colors, images, page layout) is
simplified or dropped by design.
"""

from __future__ import annotations

import re
from collections.abc import Callable
from contextlib import suppress
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Protocol

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

_MONOSPACE_FONTS = frozenset({"consolas", "courier new", "menlo", "monospace"})

_HEADING_STYLE_RE = re.compile(r"Heading (\d)", re.IGNORECASE)
_HEADING_LINE_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_LINE_RE = re.compile(r"^([-*+])\s+(.+)$")
_ORDERED_LINE_RE = re.compile(r"^(\d+)\.\s+(.+)$")
_QUOTE_LINE_RE = re.compile(r"^>\s?(.*)$")
_FENCE_LINE_RE = re.compile(r"^\s*```")
_HORIZONTAL_RULE_RE = re.compile(r"^ {0,3}([-*_]){3,}\s*$")
_PIPE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")
_SEPARATOR_CELL_RE = re.compile(r":?-+:?")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`[^`]*`|\[([^\]]*)\]\([^)]*\))")


@dataclass(frozen=True)
class DocConverter:
    """A registered document <-> Markdown converter for a set of extensions."""

    extensions: tuple[str, ...]
    language: str
    to_markdown: Callable[[bytes], str]
    from_markdown: Callable[[str], bytes]


class DocConverterProtocol(Protocol):
    """Shape of a class accepted by :func:`register_doc_converter`."""

    @staticmethod
    def to_markdown(data: bytes) -> str: ...

    @staticmethod
    def from_markdown(markdown: str) -> bytes: ...


_CONVERTERS: dict[str, DocConverter] = {}


def register_doc_converter(
    *, extensions: tuple[str, ...], language: str = "markdown"
) -> Callable[[type[DocConverterProtocol]], type[DocConverterProtocol]]:
    """Class decorator registering a converter for *extensions*.

    The decorated class must expose static ``to_markdown(bytes) -> str`` and
    ``from_markdown(str) -> bytes``.
    """

    def decorate(cls: type[DocConverterProtocol]) -> type[DocConverterProtocol]:
        converter = DocConverter(
            extensions=tuple(ext.lower() for ext in extensions),
            language=language,
            to_markdown=cls.to_markdown,
            from_markdown=cls.from_markdown,
        )
        for ext in converter.extensions:
            _CONVERTERS[ext] = converter
        return cls

    return decorate


def get_doc_converter(path: str) -> DocConverter | None:
    """Return the converter for *path*'s extension, or ``None`` if not editable."""
    return _CONVERTERS.get(_extension_of(path))


def is_editable_doc(path: str) -> bool:
    """Whether *path* has a registered editable-document extension."""
    return _extension_of(path) in _CONVERTERS


def editable_doc_extensions() -> list[str]:
    """Registered editable extensions (sorted, lowercase)."""
    return sorted(_CONVERTERS)


def _extension_of(path: str) -> str:
    name = path.replace("\\", "/").rsplit("/", 1)[-1]
    dot = name.rfind(".")
    return name[dot + 1 :].lower() if dot >= 0 else ""


# --- docx converter -------------------------------------------------------


@register_doc_converter(extensions=("docx",))
class DocxDocConverter:
    @staticmethod
    def to_markdown(data: bytes) -> str:
        # A 0-byte (or whitespace-only) file is a common artifact of creating a
        # file in the workspace UI; treat it as an empty document rather than
        # failing to parse it as a ZIP package. Saving writes back a real docx.
        if not data.strip():
            return ""
        document = Document(BytesIO(data))
        parts: list[str] = []
        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                parts.append(_paragraph_to_markdown(Paragraph(child, document)))
            elif child.tag == qn("w:tbl"):
                parts.append(_table_to_markdown(Table(child, document)))
        return "\n".join(part for part in parts if part)

    @staticmethod
    def from_markdown(markdown: str) -> bytes:
        document = Document()
        _add_markdown_blocks(document, markdown)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()


def _paragraph_to_markdown(paragraph: Paragraph) -> str:
    text = _render_paragraph(paragraph)
    prefix = _style_prefix(paragraph)
    if not text and not prefix:
        return ""
    return f"{prefix}{text}"


def _style_prefix(paragraph: Paragraph) -> str:
    style = getattr(paragraph.style, "name", None) or ""
    if style == "Title":
        return "# "
    match = _HEADING_STYLE_RE.match(style)
    if match:
        level = int(match.group(1))
        if 1 <= level <= 6:
            return "#" * level + " "
    if style == "List Bullet":
        return "- "
    if style == "List Number":
        return "1. "
    if style in ("Quote", "Intense Quote"):
        return "> "
    return ""


def _render_paragraph(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        font_name = (getattr(run.font, "name", None) or "").strip().lower()
        if run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        elif font_name in _MONOSPACE_FONTS:
            parts.append(f"`{text}`")
        else:
            parts.append(text)
    for link in getattr(paragraph, "hyperlinks", None) or []:
        if link.text:
            parts.append(link.text)
    return "".join(parts)


def _table_to_markdown(table: Table) -> str:
    rows = [[_cell_text(cell) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    lines: list[str] = []
    header = _pad_row(rows[0], width)
    lines.append("| " + " | ".join(_escape_cell(cell) for cell in header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        cells = _pad_row(row, width)
        lines.append("| " + " | ".join(_escape_cell(cell) for cell in cells) + " |")
    return "\n".join(lines)


def _pad_row(row: list[str], width: int) -> list[str]:
    return (row + [""] * width)[:width]


def _cell_text(cell: object) -> str:
    text = str(getattr(cell, "text", "") or "")
    return " ".join(text.split())


def _escape_cell(text: str) -> str:
    return text.replace("|", "\\|")


def _add_markdown_blocks(document: Any, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if _FENCE_LINE_RE.match(line):
            index += 1
            block: list[str] = []
            while index < len(lines) and not _FENCE_LINE_RE.match(lines[index]):
                block.append(lines[index])
                index += 1
            index += 1  # skip closing fence
            _add_code_block(document, block)
            continue
        if _HORIZONTAL_RULE_RE.match(line):
            index += 1
            continue
        heading = _HEADING_LINE_RE.match(line)
        if heading:
            paragraph = document.add_heading("", level=len(heading.group(1)))
            _apply_inline_runs(paragraph, heading.group(2))
            index += 1
            continue
        if _PIPE_LINE_RE.match(line):
            rows: list[str] = []
            while index < len(lines) and _PIPE_LINE_RE.match(lines[index]):
                rows.append(lines[index])
                index += 1
            _add_pipe_table(document, rows)
            continue
        bullet = _BULLET_LINE_RE.match(line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _apply_inline_runs(paragraph, bullet.group(2))
            index += 1
            continue
        ordered = _ORDERED_LINE_RE.match(line)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            _apply_inline_runs(paragraph, ordered.group(2))
            index += 1
            continue
        quote = _QUOTE_LINE_RE.match(line)
        if quote:
            paragraph = document.add_paragraph(style="Intense Quote")
            _apply_inline_runs(paragraph, quote.group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        _apply_inline_runs(paragraph, line)
        index += 1


def _add_pipe_table(document: Any, lines: list[str]) -> None:
    parsed: list[list[str]] = []
    for raw in lines:
        cells = _parse_pipe_row(raw)
        if cells is None or _is_separator_row(cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    width = max(len(row) for row in parsed)
    table = document.add_table(rows=len(parsed), cols=width)
    with suppress(Exception):
        table.style = "Table Grid"
    for row_index, row in enumerate(parsed):
        for col_index, text in enumerate(_pad_row(row, width)):
            table.rows[row_index].cells[col_index].text = text


def _parse_pipe_row(raw: str) -> list[str] | None:
    line = raw.strip()
    if not line.startswith("|") or not line.endswith("|"):
        return None
    cells: list[str] = []
    current: list[str] = []
    for char in line[1:-1]:
        if char == "\\":
            current.append(char)
            continue
        if char == "|":
            cells.append("".join(current).strip().replace("\\|", "|"))
            current = []
        else:
            current.append(char)
    cells.append("".join(current).strip().replace("\\|", "|"))
    return cells


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(_SEPARATOR_CELL_RE.fullmatch(cell) for cell in cells)


def _add_code_block(document: Any, block: list[str]) -> None:
    for line in block:
        run = document.add_paragraph().add_run(line)
        run.font.name = "Consolas"


def _apply_inline_runs(paragraph: Paragraph, text: str) -> None:
    position = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("["):
            paragraph.add_run(match.group(2) or "")
        else:
            paragraph.add_run(token[1:-1]).italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])
