"""Tests for the editable-document registry and its Markdown round-trip.

Covers :mod:`octop.infra.utils.doc_edit`: the extension registry that powers
the generic ``/workspace/doc`` endpoints, plus the docx converter's
``to_markdown`` / ``from_markdown`` round-trip.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any
from zipfile import BadZipFile

import pytest
from docx import Document

from octop.infra.utils.doc_edit import (
    DocxDocConverter,
    editable_doc_extensions,
    get_doc_converter,
    is_editable_doc,
)


def _docx_bytes(document: Any) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_sample_docx() -> bytes:
    document = Document()
    document.add_heading("Report Title", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello").bold = True
    paragraph.add_run(" world")
    document.add_paragraph("Listed item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "value"
    table.rows[1].cells[0].text = "a"
    table.rows[1].cells[1].text = "1"
    return _docx_bytes(document)


def test_registry_lists_docx() -> None:
    assert editable_doc_extensions() == ["docx"]


def test_is_editable_doc_by_extension() -> None:
    assert is_editable_doc("report.docx")
    assert is_editable_doc("/work/dir/sub/report.docx")
    assert is_editable_doc("REPORT.DOCX")
    assert not is_editable_doc("report.doc")
    assert not is_editable_doc("report.txt")
    assert not is_editable_doc("report")


def test_get_doc_converter_dispatch() -> None:
    converter = get_doc_converter("report.docx")
    assert converter is not None
    assert converter.language == "markdown"
    assert get_doc_converter("report.txt") is None


def test_to_markdown_maps_styles_and_runs() -> None:
    markdown = DocxDocConverter.to_markdown(_build_sample_docx())
    assert "# Report Title" in markdown
    assert "**Hello** world" in markdown
    assert "- Listed item" in markdown
    assert "| name | value |" in markdown
    assert "| --- | --- |" in markdown


def test_from_markdown_builds_docx() -> None:
    markdown = (
        "# Head One\n\n"
        "Some **bold** and *italic* text\n\n"
        "- Item one\n\n"
        "| a | b |\n| --- | --- |\n| 1 | 2 |\n"
    )
    data = DocxDocConverter.from_markdown(markdown)
    document = Document(BytesIO(data))
    assert [p.text for p in document.paragraphs if p.text] == [
        "Head One",
        "Some bold and italic text",
        "Item one",
    ]
    styles = [p.style.name if p.style else "" for p in document.paragraphs if p.text]
    assert styles[:2] == ["Heading 1", "Normal"]
    assert "List Bullet" in styles
    assert any(run.bold for p in document.paragraphs for run in p.runs)
    assert any(run.italic for p in document.paragraphs for run in p.runs)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.rows[1].cells[1].text == "2"


def test_round_trip_preserves_structure() -> None:
    markdown = "# Report Title\n\nHello **bold** and *italic*\n\n- Item one\n"
    markdown = DocxDocConverter.to_markdown(DocxDocConverter.from_markdown(markdown))
    assert "# Report Title" in markdown
    assert "**bold**" in markdown
    assert "*italic*" in markdown
    assert "- Item one" in markdown


def test_to_markdown_empty_file_is_empty_document() -> None:
    assert DocxDocConverter.to_markdown(b"") == ""
    assert DocxDocConverter.to_markdown(b"   \n\t ") == ""
    # An empty Markdown round-trips into a valid (empty) docx.
    data = DocxDocConverter.from_markdown("")
    document = Document(BytesIO(data))
    assert len(document.paragraphs) == 0


def test_to_markdown_rejects_non_zip() -> None:
    with pytest.raises(BadZipFile):
        DocxDocConverter.to_markdown(b"this is not a docx zip")
